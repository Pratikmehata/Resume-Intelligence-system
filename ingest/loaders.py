# ingest/loaders.py - USING VERIFIED IMPORTS
import os
from typing import List
from langchain_core.documents import Document
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pandas as pd

class ResumeLoader:
    def __init__(self, data_path: str = "./data/raw"):
        self.data_path = data_path
    
    def load_documents(self) -> List[Document]:
        """Load all documents from the data directory"""
        documents = []
        
        if not os.path.exists(self.data_path):
            os.makedirs(self.data_path, exist_ok=True)
            return documents
        
        for filename in os.listdir(self.data_path):
            file_path = os.path.join(self.data_path, filename)
            if filename.endswith('.pdf'):
                documents.extend(self._load_pdf(file_path))
            elif filename.endswith('.docx'):
                documents.extend(self._load_docx(file_path))
            elif filename.endswith('.txt'):
                documents.extend(self._load_txt(file_path))
            elif filename.endswith('.csv'):
                documents.extend(self._load_csv(file_path))
        
        return documents
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load PDF files"""
        documents = []
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "type": "resume"}
                ))
        except Exception as e:
            print(f"Error loading PDF {file_path}: {e}")
        return documents
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """Load DOCX files"""
        documents = []
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "type": "resume"}
                ))
        except Exception as e:
            print(f"Error loading DOCX {file_path}: {e}")
        return documents
    
    def _load_txt(self, file_path: str) -> List[Document]:
        """Load text files"""
        documents = []
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "type": "resume"}
                ))
        except Exception as e:
            print(f"Error loading TXT {file_path}: {e}")
        return documents
    
    def _load_csv(self, file_path: str) -> List[Document]:
        """Load CSV files (for job descriptions)"""
        documents = []
        try:
            df = pd.read_csv(file_path)
            for idx, row in df.iterrows():
                text = f"Job Title: {row.get('title', '')}\n"
                text += f"Company: {row.get('company', '')}\n"
                text += f"Description: {row.get('description', '')}\n"
                text += f"Requirements: {row.get('requirements', '')}"
                
                documents.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "type": "job_posting", "row": idx}
                ))
        except Exception as e:
            print(f"Error loading CSV {file_path}: {e}")
        return documents